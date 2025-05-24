import streamlit as st

# Set page config as the first Streamlit command
st.set_page_config(layout="wide") # Changed from "wide" to "centered"

import requests
import base64
import json
import re
import os
from dotenv import find_dotenv, load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document # For creating .docx files
from docx.shared import Pt, RGBColor # For font size and color
from docx.enum.text import WD_ALIGN_PARAGRAPH # For aligning horizontal rule
from docx.oxml.ns import qn # For qualified names in OOXML
from docx.oxml import OxmlElement # For creating OOXML elements
from io import BytesIO # To handle document in memory for download

# --- Configuration Constants (mostly for Azure DevOps) ---
# These are still "hardcoded" in the script but are now clearly defined constants.
# For deployment (e.g., Streamlit Cloud), AZURE_DEVOPS_PAT_ENV will be read from secrets.
AZURE_DEVOPS_ORG_CONFIG = "landmarkgroup"
AZURE_DEVOPS_PROJECT_CONFIG = "S2"
AZURE_API_VERSION_CONFIG = "7.2-preview"
# Set to the field name provided by the user.
# IMPORTANT: This should be the *API Reference Name* from Azure DevOps,
# not just the display name. If "Current problem statement" is only the display name,
# fetching this field might fail or return N/A.
PROBLEM_IMPACT_FIELD_NAME_CONFIG = "Current problem statement"


# --- Environment and LLM Setup ---
try:
    dotenv_path = find_dotenv(raise_error_if_not_found=False)
    if dotenv_path:
        load_dotenv(dotenv_path)
    # No sidebar warning if .env not found, as sidebar is removed.
    # else:
    #     st.sidebar.warning("`.env` file not found. LITELLM_API_KEY should be set as an environment variable.")

    LITELLM_API_KEY = os.getenv("LITELLM_API_KEY")
    AZURE_DEVOPS_PAT_ENV = os.getenv("AZURE_DEVOPS_PAT") # Load PAT from environment

    if not LITELLM_API_KEY:
        st.error("LITELLM_API_KEY environment variable not set. Please set it in your .env file or Streamlit secrets. The app cannot function without it.")
        st.stop()
    # AZURE_DEVOPS_PAT_ENV can be None if not set, will be handled by UI fallback or error later

    # Cache the LLM client
    @st.cache_resource
    def get_llm_client_cached():
        if not LITELLM_API_KEY:
            st.error("LITELLM_API_KEY is missing, cannot initialize LLM client.")
            return None
        try:
            client = ChatOpenAI(
                openai_api_base="https://lmlitellm.landmarkgroup.com", # Ensure this is correct
                default_headers={
                    "Authorization": f"Bearer {LITELLM_API_KEY}",
                    "Content-Type": "application/json"
                },
                # model="gemini-2.5-pro",
                model="landmark-gpt-4.1",
                api_key=LITELLM_API_KEY,
                temperature=0.7
            )
            return client
        except Exception as e:
            st.error(f"Failed to initialize LLM client: {e}")
            return None

    llm_client = get_llm_client_cached() # Use the cached version
    if not llm_client:
        st.stop()

except Exception as e:
    st.error(f"Error during initial setup: {e}")
    st.stop()


# --- Core Logic Functions (now use hardcoded config) ---

def extract_work_item_id(input_string):
    match = re.search(r'/(\d+)/?$', input_string)
    if match:
        return match.group(1)
    elif input_string.isdigit():
        return input_string
    return None

# Updated to accept PAT as an argument
def get_azure_devops_story_details(work_item_id, pat_to_use, progress_bar_slot):
    if not pat_to_use:
        st.error("Azure DevOps PAT was not provided or found.")
        return None

    url = f"https://dev.azure.com/{AZURE_DEVOPS_ORG_CONFIG}/{AZURE_DEVOPS_PROJECT_CONFIG}/_apis/wit/workitems/{work_item_id}?$expand=all&api-version={AZURE_API_VERSION_CONFIG}"
    
    try:
        credentials = f":{pat_to_use}"
        encoded_credentials = base64.b64encode(credentials.encode()).decode()
        headers = {"Authorization": f"Basic {encoded_credentials}", "Content-Type": "application/json"}
        
        progress_bar_slot.info(f"Fetching details for work item ID: {work_item_id}...")
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        data = response.json()
        
        details = {
            "id": data.get("id"),
            "title": data.get("fields", {}).get("System.Title", "N/A"),
            "description": data.get("fields", {}).get("System.Description", "N/A"),
            "acceptance_criteria": data.get("fields", {}).get("Microsoft.VSTS.Common.AcceptanceCriteria", "N/A"),
            "problem_impact": data.get("fields", {}).get(PROBLEM_IMPACT_FIELD_NAME_CONFIG, "N/A") # Uses hardcoded field name
        }
        progress_bar_slot.success(f"Successfully fetched details for ID: {work_item_id}")
        return details
        
    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error for ID {work_item_id}: {http_err}. Response: {response.text}")
    except requests.exceptions.RequestException as req_err:
        st.error(f"Request error for ID {work_item_id}: {req_err}")
    except json.JSONDecodeError:
        st.error(f"Failed to decode JSON for ID {work_item_id}. Response: {response.text}")
    except Exception as e:
        st.error(f"An unexpected error occurred for ID {work_item_id}: {e}")
    return None

def generate_combined_functional_document(all_stories_data, llm_instance, progress_bar_slot):
    if not all_stories_data:
        return "No story data provided to generate a document."

    def clean_text(text_content):
        if text_content is None: return "N/A"
        text_content = re.sub('<[^>]+>', '', text_content)
        html_entities = {"&nbsp;": " ", "<": "<", ">": ">", "&": "&"}
        for entity, char in html_entities.items():
            text_content = text_content.replace(entity, char)
        return text_content.strip()

    stories_summary_parts = []
    fallback_content_parts = ["Combined Functional Document (LLM Failed - Basic Fallback)\n==========================================================\n"]

    for i, details in enumerate(all_stories_data):
        story_id = details.get('id', 'N/A')
        story_title = details.get('title', 'N/A')
        description_text = clean_text(details.get('description'))
        acceptance_criteria_text = clean_text(details.get('acceptance_criteria'))
        problem_impact_text = clean_text(details.get('problem_impact'))

        stories_summary_parts.append(f"---\nStory {i+1} (ID: {story_id})\nTitle: \"{story_title}\"\nProblem/Impact: {problem_impact_text}\nDescription: {description_text}\nAcceptance Criteria: {acceptance_criteria_text}\n---")
        fallback_content_parts.extend([f"\n--- Details for Story ID: {story_id} ---\n", f"Title: {story_title}\n", f"Problem/Impact: {problem_impact_text}\n", f"Description: {description_text}\n", f"Acceptance Criteria: {acceptance_criteria_text}\n"])

    consolidated_stories_info = "\n".join(stories_summary_parts)
    combined_fallback_text = "".join(fallback_content_parts)

    prompt = f"""
You are an expert technical writer and business analyst. Your task is to create ONE consolidated, comprehensive Functional Specification Document (FSD) by synthesizing information from the provided Azure DevOps user stories.

**Overall Output Requirements:**
1.  **Document Title (First Line):** Begin your response *directly* with a concise and descriptive title for the FSD. Example: "Functional Specification: Enhanced Warehouse Short-Pick Handling System". Do NOT include any other introductory phrases on this first line.
2.  **Document Body (Following Lines):** Structure the FSD according to the sections and numbering provided below. Use Markdown headings (`#` for Level 1, `##` for Level 2, `###` for Level 3).

**FSD Structure to Follow:**

# 1. Introduction
## 1.1 Purpose of this Document
   (Describe the purpose of this FSD, focusing on what the document aims to achieve and communicate regarding the system's functionality.)
## 1.2 Project Scope and Objectives
   (Define the scope of the project/system covered by this FSD and its high-level objectives. Synthesize this from the overall goal of the user stories provided.)

# 2. System Features
   (This section details the specific functional requirements. For each distinct feature synthesized from the user stories, create a subsection. Example: If a story is about "Short Pick Handling", it becomes a system feature here. Ensure sequential numbering for features, e.g., 2.1, 2.2, etc.)

## 2.x [System Feature Name 1]
   (Replace 'x' with sequential numbering, e.g., 2.1. Replace `[System Feature Name 1]` with a descriptive name for the feature derived from the user story.)
### 2.x.1 Description and Priority
      (Provide a brief, clear description of the feature. State its priority if discernible from the input or if it can be reasonably inferred as High, Medium, or Low.)
### 2.x.2 Context and Rationale
      (Summarize the background, the core problem this feature addresses, and its business impact or the reasons it's needed. This should be based on the "Problem/Impact" information from the user stories.)
### 2.x.3 Functional Requirements
      (Detail the specific functional behaviors of this feature. This is where the core "Description" from the user stories will be elaborated into specific, testable requirements. Use bullet points or numbered lists for clarity if detailing multiple requirements within this feature.)
### 2.x.4 Verification Criteria
      (List key criteria that will be used to confirm the feature is implemented correctly and meets the requirements. This should be based on the "Acceptance Criteria" from the user stories, rephrased as testable statements.)

## 2.y [System Feature Name 2]
   (Continue this structure for all distinct features identified from the user stories. Replace 'y' with the next sequential number, e.g., 2.2)
### 2.y.1 Description and Priority
### 2.y.2 Context and Rationale
### 2.y.3 Functional Requirements
### 2.y.4 Verification Criteria

**(Continue for additional features as needed, following the 2.z numbering and sub-section structure.)**

**Important Style Guidelines:**
*   Maintain a professional and formal tone throughout the document.
*   The goal is to produce a document that reads like a standard, well-written FSD.
*   **Avoid any meta-commentary about the document itself or internal team communications.** Focus solely on describing the system and its functionalities.
*   Ensure the numbering of sections and subsections is consistent with the simplified structure provided above.

User Stories Data:
{consolidated_stories_info}
    """
    try:
        progress_bar_slot.info("Invoking LLM to generate combined document... (this may take a moment)")
        response = llm_instance.invoke(prompt)
        llm_generated_content = response.content if hasattr(response, 'content') else str(response)
        progress_bar_slot.success("LLM processing complete.")
        return llm_generated_content
    except Exception as e:
        st.error(f"Error invoking LLM: {e}")
        return combined_fallback_text

# --- Streamlit App UI ---
st.title("Functional Document Generator")
st.markdown("<p style='font-size: 16px; color: #333333;'>Generate professional Functional Specification Documents from Azure DevOps work items.</p>", unsafe_allow_html=True)
st.divider()

# Get PAT: Try from environment (for deployed app), fallback to UI input for local dev
azure_pat_from_env = os.getenv("AZURE_DEVOPS_PAT")
pat_input_field = azure_pat_from_env # Default to env var

if not azure_pat_from_env: # If not in env, show input field
    st.subheader("Azure DevOps Configuration")
    st.warning("Azure DevOps PAT not found in environment. Please enter it below for this session.")
    pat_input_field = st.text_input(
        "Personal Access Token (PAT)", 
        type="password", 
        help="Your Azure DevOps PAT with read access to work items."
    )
    if not pat_input_field:
        st.info("A PAT is required to fetch work item details from Azure DevOps.")
    st.divider()

# Main Area for Input and Output
st.subheader("Document Generation Input")
story_ids_input = st.text_area(
    label=f"Enter Azure DevOps work item links or IDs (comma-separated)",
    value="", # Ensure it starts empty or with a sensible default
    height=100, 
    placeholder=f"e.g., 12345, https://dev.azure.com/{AZURE_DEVOPS_ORG_CONFIG}/{AZURE_DEVOPS_PROJECT_CONFIG}/_workitems/edit/67890"
)

if 'generated_document' not in st.session_state:
    st.session_state.generated_document = ""
if 'processed_ids_for_run' not in st.session_state: 
    st.session_state.processed_ids_for_run = []

status_placeholder = st.empty() # For dynamic status updates

# Ensure pat_input_field is used in the button logic
if st.button("Generate Document", type="primary", use_container_width=True, key="generate_doc_button"):
    # Use pat_input_field which holds either the env var or the text input value
    current_pat_to_use = pat_input_field 
    
    if not current_pat_to_use:
        st.error("Azure DevOps PAT is required. Please set the AZURE_DEVOPS_PAT environment variable or enter it above.")
    elif not story_ids_input:
        st.error("Please enter at least one story link or ID.")
    elif not PROBLEM_IMPACT_FIELD_NAME_CONFIG:
         st.error("The 'Problem/Impact Field API Name' (PROBLEM_IMPACT_FIELD_NAME_CONFIG) is empty in app.py. Please contact the app administrator.")
    else:
        st.session_state.generated_document = ""
        st.session_state.processed_ids_for_run = []
        status_placeholder.info("Starting document generation process...")
        
        input_items = [item.strip() for item in story_ids_input.split(',') if item.strip()]
        all_stories_data = []
        has_errors = False

        if not input_items:
            st.error("No valid story links or IDs provided after parsing.")
        else:
            for item_val in input_items:
                work_item_id = extract_work_item_id(item_val)
                if work_item_id:
                    story_data = get_azure_devops_story_details(
                        work_item_id, current_pat_to_use, status_placeholder
                    )
                    if story_data:
                        all_stories_data.append(story_data)
                        st.session_state.processed_ids_for_run.append(str(work_item_id))
                    else:
                        has_errors = True 
                else:
                    status_placeholder.warning(f"Skipping invalid input: '{item_val}'")
            
            if all_stories_data and not has_errors:
                status_placeholder.info(f"All data fetched for IDs: {', '.join(st.session_state.processed_ids_for_run)}. Generating combined document...")
                combined_doc = generate_combined_functional_document(all_stories_data, llm_client, status_placeholder)
                st.session_state.generated_document = combined_doc
                if "LLM Failed" not in combined_doc:
                     status_placeholder.success("Document generation complete!")
            elif not all_stories_data:
                 status_placeholder.error("No valid story data could be fetched. Cannot generate document.")
            else:
                 status_placeholder.warning("Document generation skipped or incomplete due to errors in fetching some work items. Please review messages above.")

if st.session_state.generated_document:
    with st.expander("View Generated Functional Document", expanded=True):
        # st.header("Generated Functional Document") # Optional: Header inside expander
        st.markdown(st.session_state.generated_document)
        st.divider() # Visual separation before download buttons
    
        # Create a unique filename for download
    # Use st.session_state.processed_ids_for_run which is now correctly populated
    download_filename_suffix = "_".join(st.session_state.processed_ids_for_run) if st.session_state.processed_ids_for_run and len(st.session_state.processed_ids_for_run) < 5 else "multiple_stories"
    
    # --- Create DOCX for download ---
    try:
        doc = Document()

        # Define and apply styles
        # Normal text style
        style_normal = doc.styles['Normal']
        font_normal = style_normal.font
        font_normal.size = Pt(13)

        # Define a blue color
        blue_color_obj = RGBColor(0x1E, 0x88, 0xE5) # For use with font.color.rgb etc.
        blue_color_hex = "1E88E5"                 # Direct hex string for OOXML

        # Heading 1 style
        style_h1 = doc.styles['Heading 1']
        font_h1 = style_h1.font
        font_h1.size = Pt(20)
        font_h1.color.rgb = blue_color_obj # Use the RGBColor object
        font_h1.bold = True 
        para_format_h1 = style_h1.paragraph_format
        para_format_h1.space_after = Pt(12) # Add 12pt space after H1

        # Heading 2 style
        style_h2 = doc.styles['Heading 2']
        font_h2 = style_h2.font
        font_h2.size = Pt(16)
        font_h2.color.rgb = blue_color_obj # Use the RGBColor object
        font_h2.bold = True
        para_format_h2 = style_h2.paragraph_format
        para_format_h2.space_after = Pt(8) # Add 8pt space after H2

        # Heading 3 style
        style_h3 = doc.styles['Heading 3']
        font_h3 = style_h3.font
        font_h3.size = Pt(14)
        font_h3.color.rgb = blue_color_obj # Use the RGBColor object
        font_h3.bold = True
        para_format_h3 = style_h3.paragraph_format
        para_format_h3.space_after = Pt(6) # Add 6pt space after H3

        full_markdown_content = st.session_state.generated_document

        # Pre-processing for the entire content (e.g., em space)
        full_markdown_content = full_markdown_content.replace('\u2003', ' ') 
        full_markdown_content = full_markdown_content.strip()

        # Extract title (first line) and the rest of the content
        content_lines = full_markdown_content.split('\n', 1)
        doc_title_raw = content_lines[0].strip()
        markdown_body = content_lines[1].strip() if len(content_lines) > 1 else ""

        # Clean the extracted title (remove potential markdown heading markers)
        doc_title_cleaned = re.sub(r'^[#\s]*', '', doc_title_raw)
        
        # Further pre-processing for the body
        markdown_body = re.sub(r'(\r\n|\r|\n){3,}', '\n\n', markdown_body) # Consolidate >2 newlines
        markdown_body = markdown_body.strip()

        # Helper function to add text with inline formatting to a paragraph object
        def add_runs_to_paragraph(paragraph, text_content):
            # Regex to split by bold/italic markers, non-greedy
            # Handles **bold**, *italic*, __bold__, _italic_
            segments = re.split(r'(\*\*(?:(?!\*\*).)*\*\*|\*(?:(?!\*).)*\*|__(?:(?!__).)*__|_(?:(?!_).)*_)', text_content)
            
            for segment in segments:
                if not segment: # Skip empty strings that can result from split
                    continue
                
                run = paragraph.add_run()
                if (segment.startswith('**') and segment.endswith('**')) or \
                   (segment.startswith('__') and segment.endswith('__')): # Bold
                    run.text = segment[2:-2]
                    run.bold = True
                elif (segment.startswith('*') and segment.endswith('*')) or \
                     (segment.startswith('_') and segment.endswith('_')): # Italic
                    run.text = segment[1:-1]
                    run.italic = True
                else: # Regular text
                    # Collapse multiple whitespace chars (incl. spaces, tabs, newlines if any within segment) to a single space,
                    # and strip leading/trailing whitespace from the segment.
                    cleaned_segment_text = re.sub(r'\s+', ' ', segment).strip()
                    if cleaned_segment_text: # Only add run if there's content after cleaning
                        run.text = cleaned_segment_text

        # --- TOC and Content Processing with Bookmarks and Hyperlinks ---
        
        # First Pass: Collect headings and assign bookmark names
        collected_headings = []
        bookmark_id_counter = 0
        temp_lines = markdown_body.split('\n')

        for temp_line_raw in temp_lines:
            temp_line = temp_line_raw.strip()
            heading_level = 0
            heading_text = ""
            marker_len = 0

            if temp_line.startswith('# '): heading_level = 1; marker_len = 2
            elif temp_line.startswith('## '): heading_level = 2; marker_len = 3
            elif temp_line.startswith('### '): heading_level = 3; marker_len = 4
            
            if heading_level > 0:
                full_heading_content = temp_line_raw[marker_len:].strip()
                # For TOC display, take text before colon if present
                toc_display_text = full_heading_content
                first_colon_idx_toc = toc_display_text.find(':')
                if first_colon_idx_toc != -1:
                    toc_display_text = toc_display_text[:first_colon_idx_toc].strip()

                bookmark_id_counter += 1
                bookmark_name = f"_toc_bookmark_{bookmark_id_counter}"
                collected_headings.append({
                    "text": full_heading_content, # Full text for rendering heading later
                    "toc_text": toc_display_text, # Potentially shortened text for TOC
                    "level": heading_level,
                    "bookmark_name": bookmark_name,
                    "original_line_raw": temp_line_raw # To match during second pass
                })

        # Add "Table of Contents" heading
        toc_main_heading = doc.add_heading("Table of Contents", level=1)
        toc_main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Create TOC entries with hyperlinks
        for item in collected_headings:
            indent = ""
            if item["level"] == 2: indent = "    "
            elif item["level"] == 3: indent = "        "
            
            p = doc.add_paragraph(style='Normal') # Use Normal style for TOC entries
            p.paragraph_format.left_indent = Pt(len(indent) * 5) # Basic indentation

            # Create hyperlink
            hyperlink = OxmlElement('w:hyperlink')
            hyperlink.set(qn('w:anchor'), item["bookmark_name"])
            
            run_element = OxmlElement('w:r')
            run_props = OxmlElement('w:rPr')
            
            # Apply Hyperlink style (blue and underlined by default in Word)
            style_el = OxmlElement('w:rStyle')
            style_el.set(qn('w:val'), 'Hyperlink')
            run_props.append(style_el)

            # Explicitly set color for TOC entry (using the defined blue_color_hex)
            color_el = OxmlElement('w:color')
            color_el.set(qn('w:val'), blue_color_hex) # Use the direct hex string
            run_props.append(color_el)
            
            # Set font size for TOC entry (e.g., 13pt = 26 half-points)
            size_el = OxmlElement('w:sz')
            size_el.set(qn('w:val'), '26') # 13pt
            run_props.append(size_el)
            
            size_cs_el = OxmlElement('w:szCs') # For complex script fonts
            size_cs_el.set(qn('w:val'), '26') # 13pt
            run_props.append(size_cs_el)
            
            run_element.append(run_props)

            text_el = OxmlElement('w:t')
            text_el.text = item["toc_text"]
            run_element.append(text_el)
            
            hyperlink.append(run_element)
            p._p.append(hyperlink)

        doc.add_page_break()

        # Add the main document title (cleaned)
        if doc_title_cleaned:
            title_p = doc.add_paragraph()
            title_p.style = doc.styles['Heading 1']
            add_runs_to_paragraph(title_p, doc_title_cleaned)

        # Second Pass: Render main content with bookmarks at headings
        paragraph_buffer = []
        heading_idx_for_bookmarking = 0 # To iterate through collected_headings

        for line_raw in temp_lines: # Iterate using the same split as the first pass
            line = line_raw.strip()

            # Check if this line is a heading we've collected
            is_processed_as_heading = False
            if heading_idx_for_bookmarking < len(collected_headings) and \
               line_raw.strip() == collected_headings[heading_idx_for_bookmarking]["original_line_raw"].strip():
                
                if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []

                current_heading_data = collected_headings[heading_idx_for_bookmarking]
                heading_level = current_heading_data["level"]
                full_heading_text_to_render = current_heading_data["text"] # Use full text for rendering
                bookmark_name = current_heading_data["bookmark_name"]
                
                # Separate heading text from potential spill-over after colon
                content_after_marker = full_heading_text_to_render 
                actual_heading_text_for_render = content_after_marker 
                spill_over_text = ""
                first_colon_index = content_after_marker.find(':')
                if first_colon_index != -1:
                    potential_heading = content_after_marker[:first_colon_index+1].strip()
                    potential_spill = content_after_marker[first_colon_index+1:].strip()
                    if potential_spill: 
                        actual_heading_text_for_render = potential_heading
                        spill_over_text = potential_spill
                
                # Add the heading paragraph
                h_paragraph = doc.add_heading(level=heading_level)
                
                # Add bookmark start
                bookmark_start_el = OxmlElement('w:bookmarkStart')
                bookmark_start_el.set(qn('w:id'), str(heading_idx_for_bookmarking)) # ID needs to be unique
                bookmark_start_el.set(qn('w:name'), bookmark_name)
                h_paragraph._p.insert(0, bookmark_start_el) # Insert at the beginning of the paragraph
                
                # Add heading text
                add_runs_to_paragraph(h_paragraph, actual_heading_text_for_render)
                
                # Add bookmark end
                bookmark_end_el = OxmlElement('w:bookmarkEnd')
                bookmark_end_el.set(qn('w:id'), str(heading_idx_for_bookmarking))
                h_paragraph._p.append(bookmark_end_el) # Append at the end of the paragraph

                if spill_over_text: paragraph_buffer = [spill_over_text]
                
                heading_idx_for_bookmarking += 1
                is_processed_as_heading = True

            if is_processed_as_heading:
                continue

            # Process non-heading lines (same as before)
            if not line:
                if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                continue
            
            if line == '---' or line == '***' or line == '___':
                if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                if line == '---': doc.add_page_break()
                else: 
                     hr_p = doc.add_paragraph()
                     hr_p.add_run("_________________________________________")
                     hr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif line.startswith('* ') or line.startswith('- ') or line.startswith('+ '):
                if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                item_text = line[2:].strip()
                p_list = doc.add_paragraph(style='ListBullet')
                add_runs_to_paragraph(p_list, item_text)
            elif re.match(r'^\d+\.\s', line):
                if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                item_text = re.sub(r'^\d+\.\s', '', line).strip()
                p_list_num = doc.add_paragraph(style='ListNumber')
                add_runs_to_paragraph(p_list_num, item_text)
            else:
                paragraph_buffer.append(line_raw)

        if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), "\n".join(paragraph_buffer))
        
        bio = BytesIO()
        doc.save(bio)
        bio.seek(0)
        
        st.download_button(
            label="Download Document as Word File (.docx)",
            data=bio,
            file_name=f"functional_doc_combined_{download_filename_suffix}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    except Exception as e:
        st.error(f"Error creating .docx file: {e}")
        # Fallback to text download if docx creation fails
        st.download_button(
            label="Download Document as Text File (DOCX failed)",
            data=st.session_state.generated_document,
            file_name=f"functional_doc_combined_{download_filename_suffix}.txt",
            mime="text/plain"
        )

st.markdown("---")
