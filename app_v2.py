import streamlit as st

# Set page config as the first Streamlit command
st.set_page_config(layout="wide") # Changed from "wide" to "centered"

import requests
import base64
import json
import re
import os
from dotenv import find_dotenv, load_dotenv
from langchain_openai import ChatOpenAI
from docx import Document # For creating .docx files
from docx.shared import Pt, RGBColor # For font size and color
from docx.enum.text import WD_ALIGN_PARAGRAPH # For aligning horizontal rule
from docx.oxml.ns import qn # For qualified names in OOXML
from docx.oxml import OxmlElement # For creating OOXML elements
from io import BytesIO # To handle document in memory for download

# --- Configuration Constants (mostly for Azure DevOps) ---
# Set to the field name provided by the user.
# IMPORTANT: This should be the *API Reference Name* from Azure DevOps,
# not just the display name. If "Current problem statement" is only the display name,
# fetching this field might fail or return N/A.
PROBLEM_IMPACT_FIELD_NAME_CONFIG = "Current problem statement"


# --- Environment and LLM Setup ---
try:
    dotenv_path = find_dotenv(raise_error_if_not_found=False)
    if dotenv_path:
        load_dotenv(dotenv_path)
    # No sidebar warning if .env not found, as sidebar is removed.
    # else:
    #     st.sidebar.warning("`.env` file not found. LITELLM_API_KEY should be set as an environment variable.")

    LITELLM_API_KEY = os.getenv("LITELLM_API_KEY")
    AZURE_DEVOPS_PAT_ENV = os.getenv("AZURE_DEVOPS_PAT") # Load PAT from environment

    # Load Azure DevOps config *after* dotenv is loaded
    AZURE_DEVOPS_ORG_CONFIG = os.getenv("AZURE_DEVOPS_ORG_CONFIG")
    AZURE_DEVOPS_PROJECT_CONFIG = os.getenv("AZURE_DEVOPS_PROJECT_CONFIG")
    AZURE_API_VERSION_CONFIG = os.getenv("AZURE_API_VERSION_CONFIG")

    if not LITELLM_API_KEY:
        st.error("LITELLM_API_KEY environment variable not set. Please set it in your .env file or Streamlit secrets. The app cannot function without it.")
        st.stop()
    # AZURE_DEVOPS_PAT_ENV can be None if not set, will be handled by UI fallback or error later

    # Cache the LLM client
    @st.cache_resource
    def get_llm_client_cached():
        if not LITELLM_API_KEY:
            st.error("LITELLM_API_KEY is missing, cannot initialize LLM client.")
            return None
        try:
            client = ChatOpenAI(
                openai_api_base="https://lmlitellm.landmarkgroup.com", # Ensure this is correct
                default_headers={
                    "Authorization": f"Bearer {LITELLM_API_KEY}",
                    "Content-Type": "application/json"
                },
                # model="gemini-2.5-pro",
                model="landmark-gpt-4.1",
                api_key=LITELLM_API_KEY,
                temperature=0.7
            )
            return client
        except Exception as e:
            st.error(f"Failed to initialize LLM client: {e}")
            return None

    llm_client = get_llm_client_cached() # Use the cached version
    if not llm_client:
        st.stop()

except Exception as e:
    st.error(f"Error during initial setup: {e}")
    st.stop()


# --- Core Logic Functions (now use hardcoded config) ---

def get_child_story_ids(parent_work_item_id, pat_to_use, progress_bar_slot):
    """
    Fetches all child work item IDs of type 'User Story' for a given parent work item.
    """
    if not pat_to_use:
        st.error("Azure DevOps PAT was not provided or found.")
        return []

    url = f"https://dev.azure.com/{AZURE_DEVOPS_ORG_CONFIG}/{AZURE_DEVOPS_PROJECT_CONFIG}/_apis/wit/workitems/{parent_work_item_id}?$expand=relations&api-version={AZURE_API_VERSION_CONFIG}"
    try:
        credentials = f":{pat_to_use}"
        encoded_credentials = base64.b64encode(credentials.encode()).decode()
        headers = {"Authorization": f"Basic {encoded_credentials}", "Content-Type": "application/json"}

        progress_bar_slot.info(f"Fetching child work items for parent ID: {parent_work_item_id}...")
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        data = response.json()

        child_ids = []
        relations = data.get("relations", [])
        for rel in relations:
            if rel.get("rel") == "System.LinkTypes.Hierarchy-Forward":
                url = rel.get("url", "")
                match = re.search(r'/workItems/(\d+)$', url)
                if match:
                    child_ids.append(match.group(1))

        # Now, filter only "User Story" work items
        story_ids = []
        if child_ids:
            # Azure DevOps batch API: up to 200 IDs per call
            batch_size = 200
            for i in range(0, len(child_ids), batch_size):
                batch = child_ids[i:i+batch_size]
                batch_ids_str = ",".join(batch)
                batch_url = f"https://dev.azure.com/{AZURE_DEVOPS_ORG_CONFIG}/{AZURE_DEVOPS_PROJECT_CONFIG}/_apis/wit/workitemsbatch?api-version={AZURE_API_VERSION_CONFIG}"
                payload = {
                    "ids": [int(cid) for cid in batch],
                    "fields": ["System.WorkItemType"]
                }
                batch_resp = requests.post(batch_url, headers=headers, json=payload, timeout=30)
                batch_resp.raise_for_status()
                batch_data = batch_resp.json()
                for item in batch_data.get("value", []):
                    if item.get("fields", {}).get("System.WorkItemType", "").lower() == "user story":
                        story_ids.append(str(item.get("id")))
        progress_bar_slot.success(f"Found {len(story_ids)} child user stories.")
        return story_ids

    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error fetching children for parent ID {parent_work_item_id}: {http_err}. Response: {response.text}")
    except requests.exceptions.RequestException as req_err:
        st.error(f"Request error fetching children for parent ID {parent_work_item_id}: {req_err}")
    except Exception as e:
        st.error(f"An unexpected error occurred fetching children for parent ID {parent_work_item_id}: {e}")
    return []

def extract_work_item_id(input_string):
    match = re.search(r'/(\d+)/?$', input_string)
    if match:
        return match.group(1)
    elif input_string.isdigit():
        return input_string
    return None

# Updated to accept PAT as an argument
def get_azure_devops_story_details(work_item_id, pat_to_use, progress_bar_slot):
    if not pat_to_use:
        st.error("Azure DevOps PAT was not provided or found.")
        return None

    url = f"https://dev.azure.com/{AZURE_DEVOPS_ORG_CONFIG}/{AZURE_DEVOPS_PROJECT_CONFIG}/_apis/wit/workitems/{work_item_id}?$expand=all&api-version={AZURE_API_VERSION_CONFIG}"
    
    try:
        credentials = f":{pat_to_use}"
        encoded_credentials = base64.b64encode(credentials.encode()).decode()
        headers = {"Authorization": f"Basic {encoded_credentials}", "Content-Type": "application/json"}
        
        progress_bar_slot.info(f"Fetching details for work item ID: {work_item_id}...")
        response = requests.get(url, headers=headers, timeout=30)
        response.raise_for_status()
        data = response.json()
        
        details = {
            "id": data.get("id"),
            "title": data.get("fields", {}).get("System.Title", "N/A"),
            "description": data.get("fields", {}).get("System.Description", "N/A"),
            "acceptance_criteria": data.get("fields", {}).get("Microsoft.VSTS.Common.AcceptanceCriteria", "N/A"),
            "problem_impact": data.get("fields", {}).get(PROBLEM_IMPACT_FIELD_NAME_CONFIG, "N/A") # Uses hardcoded field name
        }
        progress_bar_slot.success(f"Successfully fetched details for ID: {work_item_id}")
        return details
        
    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error for ID {work_item_id}: {http_err}. Response: {response.text}")
    except requests.exceptions.RequestException as req_err:
        st.error(f"Request error for ID {work_item_id}: {req_err}")
    except json.JSONDecodeError:
        st.error(f"Failed to decode JSON for ID {work_item_id}. Response: {response.text}")
    except Exception as e:
        st.error(f"An unexpected error occurred for ID {work_item_id}: {e}")
    return None

def generate_combined_functional_document(all_stories_data, llm_instance, progress_bar_slot):
    if not all_stories_data:
        return "No story data provided to generate a document."

    def clean_text(text_content):
        if text_content is None: return "N/A"
        text_content = re.sub('<[^>]+>', '', text_content)
        html_entities = {"&nbsp;": " ", "<": "<", ">": ">", "&": "&"}
        for entity, char in html_entities.items():
            text_content = text_content.replace(entity, char)
        return text_content.strip()

    stories_summary_parts = []
    fallback_content_parts = ["Combined Functional Document (LLM Failed - Basic Fallback)\n==========================================================\n"]

    for i, details in enumerate(all_stories_data):
        story_id = details.get('id', 'N/A')
        story_title = details.get('title', 'N/A')
        description_text = clean_text(details.get('description'))
        acceptance_criteria_text = clean_text(details.get('acceptance_criteria'))
        problem_impact_text = clean_text(details.get('problem_impact'))

        stories_summary_parts.append(f"---\nStory {i+1} (ID: {story_id})\nTitle: \"{story_title}\"\nProblem/Impact: {problem_impact_text}\nDescription: {description_text}\nAcceptance Criteria: {acceptance_criteria_text}\n---")
        fallback_content_parts.extend([f"\n--- Details for Story ID: {story_id} ---\n", f"Title: {story_title}\n", f"Problem/Impact: {problem_impact_text}\n", f"Description: {description_text}\n", f"Acceptance Criteria: {acceptance_criteria_text}\n"])

    consolidated_stories_info = "\n".join(stories_summary_parts)
    combined_fallback_text = "".join(fallback_content_parts)

    prompt = (
        "You are an expert technical writer and business analyst. Your task is to create ONE consolidated, comprehensive Functional Specification Document (FSD) by synthesizing information from the provided Azure DevOps user stories.\n\n"
        "**Overall Output Requirements:**\n"
        "1.  **Document Title (First Line):** Begin your response *directly* with a concise and descriptive title for the FSD. Example: \"Functional Specification: Enhanced Warehouse Short-Pick Handling System\". Do NOT include any other introductory phrases on this first line.\n"
        "2.  **Document Body (Following Lines):** Structure the FSD according to the sections and numbering provided below. Use Markdown headings (`#` for Level 1, `##` for Level 2, `###` for Level 3).\n"
        "3.  **Flow Diagram:** At the end of your response, include a system/process flow diagram in Mermaid format, using a ```mermaid code block. The diagram should visually represent the main process or workflow described by the user stories, showing key steps, decisions, and actors.\n\n"
        "**IMPORTANT:** Use the following as a template for your Mermaid code block to ensure correct syntax:\n"
        "~~~mermaid\n"
        "flowchart TD\n"
        "    A[Homepage: Bulk & Single Item Sorting Tiles] --> B[Select Sorting Type]\n"
        "    B --> C[Fetch Pending Waves]\n"
        "    C --> D[Select Wave & View Stores]\n"
        "    D --> E{Grid Building Needed?}\n"
        "    E -- \"Yes\" --> F[Scan or Enter Grid ID]\n"
        "    F --> G[Validate Grid ID - WMS]\n"
        "    G --> H[Carton and Grid Location Mapping]\n"
        "    H --> I[Confirm Grid Building]\n"
        "    I --> J[Store Carton-Location Mapping in WMS]\n"
        "    J --> K[Scan Source Carton]\n"
        "    E -- \"No\" --> K\n"
        "    K --> L[Validate Source Carton and Fetch Item List]\n"
        "    L --> M[Scan Item]\n"
        "    M --> N[Validate Item and Fetch Store Distribution]\n"
        "    N --> O[Sort Item to Store Cartons]\n"
        "    O --> P[Confirm Sorting]\n"
        "    P --> Q{{More Items in Carton?}}\n"
        "    Q -- \"Yes\" --> M\n"
        "    Q -- \"No\" --> R{{More Source Cartons?}}\n"
        "    R -- \"Yes\" --> K\n"
        "    R -- \"No\" --> S[Sorting Complete: Return to Wave Selection]\n"
        "    O -.-> T[Force Close / Short Pick / Skip Carton / Add Container]\n"
        "    T --> O\n"
        "~~~\n"
        "Replace the nodes and flow as needed to match the user stories, but always follow this structure for valid Mermaid syntax.\n\n"
        "**FSD Structure to Follow:**\n\n"
        "# 1. Introduction\n"
        "## 1.1 Purpose of this Document\n"
        "   (Describe the purpose of this FSD, focusing on what the document aims to achieve and communicate regarding the system's functionality.)\n"
        "## 1.2 Project Scope and Objectives\n"
        "   (Define the scope of the project/system covered by this FSD and its high-level objectives. Synthesize this from the overall goal of the user stories provided.)\n\n"
        "# 2. System Features\n"
        "   (This section details the specific functional requirements. For each distinct feature synthesized from the user stories, create a subsection. Example: If a story is about \"Short Pick Handling\", it becomes a system feature here. Ensure sequential numbering for features, e.g., 2.1, 2.2, etc.)\n\n"
        "## 2.x [System Feature Name 1]\n"
        "   (Replace 'x' with sequential numbering, e.g., 2.1. Replace `[System Feature Name 1]` with a descriptive name for the feature derived from the user story.)\n"
        "### 2.x.1 Description\n"
        "      (Provide a brief, clear description of the feature.)\n"
        "### 2.x.2 Context and Rationale\n"
        "      (Summarize the background, the core problem this feature addresses, and its business impact or the reasons it's needed. This should be based on the \"Problem/Impact\" information from the user stories.)\n"
        "### 2.x.3 Functional Requirements\n"
        "      (Detail the specific functional behaviors of this feature. This is where the core \"Description\" from the user stories will be elaborated into specific, testable requirements. Use bullet points or numbered lists for clarity if detailing multiple requirements within this feature.)\n"
        "### 2.x.4 Verification Criteria\n"
        "      (List key criteria that will be used to confirm the feature is implemented correctly and meets the requirements. This should be based on the \"Acceptance Criteria\" from the user stories, rephrased as testable statements.)\n\n"
        "## 2.y [System Feature Name 2]\n"
        "   (Continue this structure for all distinct features identified from the user stories. Replace 'y' with the next sequential number, e.g., 2.2)\n"
        "### 2.y.1 Description\n"
        "### 2.y.2 Context and Rationale\n"
        "### 2.y.3 Functional Requirements\n"
        "### 2.y.4 Verification Criteria\n\n"
        "**(Continue for additional features as needed, following the 2.z numbering and sub-section structure.)**\n\n"
        "**Important Style Guidelines:**\n"
        "*   Maintain a professional and formal tone throughout the document.\n"
        "*   The goal is to produce a document that reads like a standard, well-written FSD.\n"
        "*   **Avoid any meta-commentary about the document itself or internal team communications.** Focus solely on describing the system and its functionalities.\n"
        "*   Ensure the numbering of sections and subsections is consistent with the simplified structure provided above.\n\n"
        "User Stories Data:\n"
        f"{consolidated_stories_info}\n"
    )
    try:
        progress_bar_slot.info("Invoking LLM to generate combined document... (this may take a moment)")
        response = llm_instance.invoke(prompt)
        llm_generated_content = response.content if hasattr(response, 'content') else str(response)
        progress_bar_slot.success("LLM processing complete.")
        return llm_generated_content
    except Exception as e:
        st.error(f"Error invoking LLM: {e}")
        return combined_fallback_text

# --- Streamlit App UI ---
st.title("Functional Document Generator")
st.markdown("<p style='font-size: 16px; color: #333333;'>Generate professional Functional Specification Documents from Azure DevOps work items.</p>", unsafe_allow_html=True)
st.divider()

# Get PAT: Try from environment (for deployed app), fallback to UI input for local dev
azure_pat_from_env = os.getenv("AZURE_DEVOPS_PAT")
pat_input_field = azure_pat_from_env # Default to env var

if not azure_pat_from_env: # If not in env, show input field
    st.subheader("Azure DevOps Configuration")
    st.warning("Azure DevOps PAT not found in environment. Please enter it below for this session.")
    pat_input_field = st.text_input(
        "Personal Access Token (PAT)", 
        type="password", 
        help="Your Azure DevOps PAT with read access to work items."
    )
    if not pat_input_field:
        st.info("A PAT is required to fetch work item details from Azure DevOps.")
    st.divider()

# Main Area for Input and Output
st.subheader("Document Generation Input")
parent_id_input = st.text_input(
    label=f"Enter Azure DevOps parent work item link or ID",
    value="", # Ensure it starts empty or with a sensible default
    placeholder=f"e.g., 12345, https://dev.azure.com/{AZURE_DEVOPS_ORG_CONFIG}/{AZURE_DEVOPS_PROJECT_CONFIG}/_workitems/edit/67890"
)

if 'generated_document' not in st.session_state:
    st.session_state.generated_document = ""
if 'processed_ids_for_run' not in st.session_state: 
    st.session_state.processed_ids_for_run = []

status_placeholder = st.empty() # For dynamic status updates

# Ensure pat_input_field is used in the button logic
if st.button("Generate Document", type="primary", use_container_width=True, key="generate_doc_button"):
    # Use pat_input_field which holds either the env var or the text input value
    current_pat_to_use = pat_input_field 

    if not current_pat_to_use:
        st.error("Azure DevOps PAT is required. Please set the AZURE_DEVOPS_PAT environment variable or enter it above.")
    elif not parent_id_input:
        st.error("Please enter a parent work item link or ID.")
    elif not PROBLEM_IMPACT_FIELD_NAME_CONFIG:
        st.error("The 'Problem/Impact Field API Name' (PROBLEM_IMPACT_FIELD_NAME_CONFIG) is empty in app.py. Please contact the app administrator.")
    else:
        st.session_state.generated_document = ""
        st.session_state.processed_ids_for_run = []
        status_placeholder.info("Starting document generation process...")

        # Support multiple comma-separated IDs/links
        input_items = [x.strip() for x in parent_id_input.split(",") if x.strip()]
        if not input_items:
            st.error("Invalid work item link(s) or ID(s).")
        else:
            all_story_ids_to_fetch = set()
            parent_ids = []
            child_ids = []
            parent_to_child_ids = {}
            status_placeholder.info("Processing input work items...")

            # First pass: determine which are parents (have children) and which are normal
            for item in input_items:
                work_item_id = extract_work_item_id(item)
                if not work_item_id:
                    st.warning(f"Invalid work item link or ID: {item}")
                    continue
                # Check for child stories
                child_story_ids = get_child_story_ids(work_item_id, current_pat_to_use, status_placeholder)
                if child_story_ids:
                    parent_ids.append(work_item_id)
                    parent_to_child_ids[work_item_id] = child_story_ids
                    all_story_ids_to_fetch.add(work_item_id)
                    for cid in child_story_ids:
                        all_story_ids_to_fetch.add(cid)
                else:
                    child_ids.append(work_item_id)
                    all_story_ids_to_fetch.add(work_item_id)

            # Deduplicate
            all_story_ids_to_fetch = list(all_story_ids_to_fetch)
            st.session_state.processed_ids_for_run = []

            all_stories_data = []
            has_errors = False

            # Fetch all story details
            for sid in all_story_ids_to_fetch:
                story_data = get_azure_devops_story_details(sid, current_pat_to_use, status_placeholder)
                raw_json_payload = None
                try:
                    AZURE_DEVOPS_ORG_CONFIG = os.getenv("AZURE_DEVOPS_ORG_CONFIG")
                    AZURE_DEVOPS_PROJECT_CONFIG = os.getenv("AZURE_DEVOPS_PROJECT_CONFIG")
                    AZURE_API_VERSION_CONFIG = os.getenv("AZURE_API_VERSION_CONFIG")
                    credentials = f":{current_pat_to_use}"
                    encoded_credentials = base64.b64encode(credentials.encode()).decode()
                    headers = {"Authorization": f"Basic {encoded_credentials}", "Content-Type": "application/json"}
                    url = f"https://dev.azure.com/{AZURE_DEVOPS_ORG_CONFIG}/{AZURE_DEVOPS_PROJECT_CONFIG}/_apis/wit/workitems/{sid}?$expand=all&api-version={AZURE_API_VERSION_CONFIG}"
                    resp = requests.get(url, headers=headers, timeout=30)
                    resp.raise_for_status()
                    raw_json_payload = resp.json()
                except Exception:
                    raw_json_payload = None

                if story_data:
                    story_data["_raw_json_payload"] = raw_json_payload
                    all_stories_data.append(story_data)
                    st.session_state.processed_ids_for_run.append(str(sid))
                else:
                    has_errors = True

            if all_stories_data and not has_errors:
                status_placeholder.info(f"All data fetched for work items: {', '.join(st.session_state.processed_ids_for_run)}. Generating combined document...")
                combined_doc = generate_combined_functional_document(all_stories_data, llm_client, status_placeholder)
                st.session_state.generated_document = combined_doc
                st.session_state.all_stories_data = all_stories_data  # Store for expander
                if "LLM Failed" not in combined_doc:
                    status_placeholder.success("Document generation complete!")
            elif not all_stories_data:
                status_placeholder.error("No valid story data could be fetched. Cannot generate document.")
            else:
                status_placeholder.warning("Document generation skipped or incomplete due to errors in fetching some work items. Please review messages above.")

if st.session_state.generated_document:
    import streamlit.components.v1 as components

    with st.expander("View Generated Functional Document", expanded=True):
        # Extract and render Mermaid diagram if present
        mermaid_match = re.search(r"```mermaid\s*([\s\S]+?)```", st.session_state.generated_document)
        doc_to_display = st.session_state.generated_document
        if mermaid_match:
            # For single story, do not show diagram if it would be a single node
            if hasattr(st.session_state, "all_stories_data") and len(st.session_state.all_stories_data) == 1:
                story = st.session_state.all_stories_data[0]
                story_title = story.get("title", "User Story")
                # If the fallback diagram would be a single node, skip rendering
                # (i.e., do nothing: do not show diagram, do not show code block)
                pass
            else:
                mermaid_code = mermaid_match.group(1)
                # Clean up Mermaid code: remove leading/trailing whitespace and normalize indentation
                mermaid_code_clean = "\n".join([line.lstrip() for line in mermaid_code.strip().splitlines() if line.strip()])
                # Remove the mermaid code block from the markdown before displaying
                doc_to_display = re.sub(r"```mermaid\s*([\s\S]+?)```", "", doc_to_display)

                # --- Mermaid v11.6.0 compatibility check ---
                def is_mermaid_v11_compatible(code: str) -> (bool, str):
                    # Basic checks: must start with 'flowchart TD' and not contain unsupported syntax
                    lines = [l.strip() for l in code.splitlines() if l.strip()]
                    if not lines or not lines[0].startswith("flowchart TD"):
                        return False, "Mermaid diagram must start with 'flowchart TD' for compatibility with version 11.6.0."
                    # Check for unsupported constructs (e.g., subgraph, new syntax, etc.)
                    unsupported_keywords = ["subgraph", "click ", ":::"]  # Add more as needed
                    for kw in unsupported_keywords:
                        if any(kw in l for l in lines):
                            return False, f"Mermaid diagram contains unsupported syntax for v11.6.0: '{kw}'."
                    # Must have at least one edge (--> or --)
                    has_edge = any("-->" in l or "--" in l for l in lines[1:])
                    if not has_edge:
                        return False, "Mermaid diagram must contain at least one edge (e.g., 'A --> B')."
                    # No empty diagrams
                    if len(lines) < 2:
                        return False, "Mermaid diagram is empty."
                    # Suppress diagram if only a single node (no edges)
                    node_lines = [l for l in lines[1:] if "[" in l and "]" in l]
                    if has_edge is False and len(node_lines) == 1:
                        return False, "Mermaid diagram contains only a single node; diagram suppressed."
                    return True, ""

                is_valid, err_msg = is_mermaid_v11_compatible(mermaid_code_clean)
                import uuid
                unique_id = f"mermaid-container-{uuid.uuid4().hex}"
                if is_valid:
                    components.html(
                        f"""
                        <div id="{unique_id}" style="border:2px solid #888; min-height:400px; margin-bottom:10px; overflow:auto; max-height:1200px;">
                          <div class="mermaid">
{mermaid_code_clean}
                          </div>
                        </div>
                        <script src="https://cdn.jsdelivr.net/npm/mermaid/dist/mermaid.min.js"></script>
                        <script>
                        if (window.mermaid) {{
                            mermaid.initialize({{startOnLoad:true}});
                            mermaid.init(undefined, "#{unique_id} .mermaid");
                        }}
                        </script>
                        """,
                        height=1200,
                    )
                else:
                    # Suppress all invalid diagrams (do not show warning or code block)
                    pass
        # Show the document without the mermaid code block
        st.markdown(doc_to_display)
        st.divider() # Visual separation before download buttons

        # Show JSON message structures found in story fields (not the full payload)
        if hasattr(st.session_state, "all_stories_data"):
            any_json_found = False
            for i, story in enumerate(st.session_state.all_stories_data, 1):
                found_jsons = []
                for field in ["description", "acceptance_criteria"]:
                    val = story.get(field)
                    if isinstance(val, str):
                        matches = re.findall(r'({[\s\S]+?})', val)
                        for match in matches:
                            try:
                                parsed = json.loads(match)
                                found_jsons.append(parsed)
                            except Exception:
                                continue
                if found_jsons:
                    any_json_found = True
                    st.markdown(f"**Story {i} (ID: {story.get('id', 'N/A')}) - JSON Message Structures Found:**")
                    for j, js in enumerate(found_jsons, 1):
                        st.code(json.dumps(js, indent=2), language="json")
            if any_json_found:
                st.divider()

        # Create a unique filename for download
        download_filename_suffix = "_".join(st.session_state.processed_ids_for_run) if st.session_state.processed_ids_for_run and len(st.session_state.processed_ids_for_run) < 5 else "multiple_stories"

        # --- Create DOCX for download ---
        try:
            doc = Document()

            # Define and apply styles
            style_normal = doc.styles['Normal']
            font_normal = style_normal.font
            font_normal.size = Pt(13)

            blue_color_obj = RGBColor(0x1E, 0x88, 0xE5)
            blue_color_hex = "1E88E5"

            style_h1 = doc.styles['Heading 1']
            font_h1 = style_h1.font
            font_h1.size = Pt(20)
            font_h1.color.rgb = blue_color_obj
            font_h1.bold = True
            para_format_h1 = style_h1.paragraph_format
            para_format_h1.space_after = Pt(12)

            style_h2 = doc.styles['Heading 2']
            font_h2 = style_h2.font
            font_h2.size = Pt(16)
            font_h2.color.rgb = blue_color_obj
            font_h2.bold = True
            para_format_h2 = style_h2.paragraph_format
            para_format_h2.space_after = Pt(8)

            style_h3 = doc.styles['Heading 3']
            font_h3 = style_h3.font
            font_h3.size = Pt(14)
            font_h3.color.rgb = blue_color_obj
            font_h3.bold = True
            para_format_h3 = style_h3.paragraph_format
            para_format_h3.space_after = Pt(6)

            full_markdown_content = st.session_state.generated_document
            full_markdown_content = full_markdown_content.replace('\u2003', ' ')
            full_markdown_content = full_markdown_content.strip()

            # Remove Mermaid code block before parsing markdown
            mermaid_match_docx = re.search(r"```mermaid\s*([\s\S]+?)```", full_markdown_content)
            mermaid_code_docx_clean = None
            if mermaid_match_docx:
                mermaid_code_docx = mermaid_match_docx.group(1)
                mermaid_code_docx_clean = "\n".join([line.lstrip() for line in mermaid_code_docx.strip().splitlines() if line.strip()])
                # Remove the code block from the markdown
                full_markdown_content = re.sub(r"```mermaid\s*([\s\S]+?)```", "", full_markdown_content)

            content_lines = full_markdown_content.split('\n', 1)
            doc_title_raw = content_lines[0].strip()
            markdown_body = content_lines[1].strip() if len(content_lines) > 1 else ""

            doc_title_cleaned = re.sub(r'^[#\s]*', '', doc_title_raw)
            markdown_body = re.sub(r'(\r\n|\r|\n){3,}', '\n\n', markdown_body)
            markdown_body = markdown_body.strip()

            def add_runs_to_paragraph(paragraph, text_content):
                segments = re.split(r'(\*\*(?:(?!\*\*).)*\*\*|\*(?:(?!\*).)*\*|__(?:(?!__).)*__|_(?:(?!_).)*_)', text_content)
                for segment in segments:
                    if not segment:
                        continue
                    run = paragraph.add_run()
                    if (segment.startswith('**') and segment.endswith('**')) or \
                       (segment.startswith('__') and segment.endswith('__')):
                        run.text = segment[2:-2]
                        run.bold = True
                    elif (segment.startswith('*') and segment.endswith('*')) or \
                         (segment.startswith('_') and segment.endswith('_')):
                        run.text = segment[1:-1]
                        run.italic = True
                    else:
                        cleaned_segment_text = re.sub(r'\s+', ' ', segment).strip()
                        if cleaned_segment_text:
                            run.text = cleaned_segment_text

            # TOC and Content Processing with Bookmarks and Hyperlinks
            collected_headings = []
            bookmark_id_counter = 0
            temp_lines = markdown_body.split('\n')

            for temp_line_raw in temp_lines:
                temp_line = temp_line_raw.strip()
                heading_level = 0
                marker_len = 0
                if temp_line.startswith('# '): heading_level = 1; marker_len = 2
                elif temp_line.startswith('## '): heading_level = 2; marker_len = 3
                elif temp_line.startswith('### '): heading_level = 3; marker_len = 4
                if heading_level > 0:
                    full_heading_content = temp_line_raw[marker_len:].strip()
                    toc_display_text = full_heading_content
                    first_colon_idx_toc = toc_display_text.find(':')
                    if first_colon_idx_toc != -1:
                        toc_display_text = toc_display_text[:first_colon_idx_toc].strip()
                    bookmark_id_counter += 1
                    bookmark_name = f"_toc_bookmark_{bookmark_id_counter}"
                    collected_headings.append({
                        "text": full_heading_content,
                        "toc_text": toc_display_text,
                        "level": heading_level,
                        "bookmark_name": bookmark_name,
                        "original_line_raw": temp_line_raw
                    })

            toc_main_heading = doc.add_heading("Table of Contents", level=1)
            toc_main_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for item in collected_headings:
                indent = ""
                if item["level"] == 2: indent = "    "
                elif item["level"] == 3: indent = "        "
                p = doc.add_paragraph(style='Normal')
                p.paragraph_format.left_indent = Pt(len(indent) * 5)
                hyperlink = OxmlElement('w:hyperlink')
                hyperlink.set(qn('w:anchor'), item["bookmark_name"])
                run_element = OxmlElement('w:r')
                run_props = OxmlElement('w:rPr')
                style_el = OxmlElement('w:rStyle')
                style_el.set(qn('w:val'), 'Hyperlink')
                run_props.append(style_el)
                color_el = OxmlElement('w:color')
                color_el.set(qn('w:val'), blue_color_hex)
                run_props.append(color_el)
                size_el = OxmlElement('w:sz')
                size_el.set(qn('w:val'), '26')
                run_props.append(size_el)
                size_cs_el = OxmlElement('w:szCs')
                size_cs_el.set(qn('w:val'), '26')
                run_props.append(size_cs_el)
                run_element.append(run_props)
                text_el = OxmlElement('w:t')
                text_el.text = item["toc_text"]
                run_element.append(text_el)
                hyperlink.append(run_element)
                p._p.append(hyperlink)

            doc.add_page_break()

            if doc_title_cleaned:
                title_p = doc.add_paragraph()
                title_p.style = doc.styles['Heading 1']
                add_runs_to_paragraph(title_p, doc_title_cleaned)

            paragraph_buffer = []
            heading_idx_for_bookmarking = 0

            for line_raw in temp_lines:
                line = line_raw.strip()
                is_processed_as_heading = False
                if heading_idx_for_bookmarking < len(collected_headings) and \
                   line_raw.strip() == collected_headings[heading_idx_for_bookmarking]["original_line_raw"].strip():
                    if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                    current_heading_data = collected_headings[heading_idx_for_bookmarking]
                    heading_level = current_heading_data["level"]
                    full_heading_text_to_render = current_heading_data["text"]
                    bookmark_name = current_heading_data["bookmark_name"]
                    content_after_marker = full_heading_text_to_render
                    actual_heading_text_for_render = content_after_marker
                    spill_over_text = ""
                    first_colon_index = content_after_marker.find(':')
                    if first_colon_index != -1:
                        potential_heading = content_after_marker[:first_colon_index+1].strip()
                        potential_spill = content_after_marker[first_colon_index+1:].strip()
                        if potential_spill:
                            actual_heading_text_for_render = potential_heading
                            spill_over_text = potential_spill
                    h_paragraph = doc.add_heading(level=heading_level)
                    bookmark_start_el = OxmlElement('w:bookmarkStart')
                    bookmark_start_el.set(qn('w:id'), str(heading_idx_for_bookmarking))
                    bookmark_start_el.set(qn('w:name'), bookmark_name)
                    h_paragraph._p.insert(0, bookmark_start_el)
                    add_runs_to_paragraph(h_paragraph, actual_heading_text_for_render)
                    bookmark_end_el = OxmlElement('w:bookmarkEnd')
                    bookmark_end_el.set(qn('w:id'), str(heading_idx_for_bookmarking))
                    h_paragraph._p.append(bookmark_end_el)
                    if spill_over_text: paragraph_buffer = [spill_over_text]
                    heading_idx_for_bookmarking += 1
                    is_processed_as_heading = True
                if is_processed_as_heading:
                    continue
                if not line:
                    if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                    continue
                if line == '---' or line == '***' or line == '___':
                    if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                    if line == '---': doc.add_page_break()
                    else:
                        hr_p = doc.add_paragraph()
                        hr_p.add_run("_________________________________________")
                        hr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif line.startswith('* ') or line.startswith('- ') or line.startswith('+ '):
                    if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                    item_text = line[2:].strip()
                    p_list = doc.add_paragraph(style='ListBullet')
                    add_runs_to_paragraph(p_list, item_text)
                elif re.match(r'^\d+\.\s', line):
                    if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), " ".join(paragraph_buffer)); paragraph_buffer = []
                    item_text = re.sub(r'^\d+\.\s', '', line).strip()
                    p_list_num = doc.add_paragraph(style='ListNumber')
                    add_runs_to_paragraph(p_list_num, item_text)
                else:
                    paragraph_buffer.append(line_raw)
            if paragraph_buffer: add_runs_to_paragraph(doc.add_paragraph(), "\n".join(paragraph_buffer))

            # Insert Mermaid diagram as image if present (move to top, after title)
            if mermaid_code_docx_clean:
                try:
                    import requests
                    import tempfile
                    # Use kroki.io to render Mermaid to PNG
                    response = requests.post(
                        "https://kroki.io/mermaid/png",
                        data=mermaid_code_docx_clean.encode("utf-8"),
                        headers={"Content-Type": "text/plain"},
                        timeout=30,
                    )
                    if response.status_code == 200:
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp_img:
                            tmp_img.write(response.content)
                            tmp_img.flush()
                            # Insert after title, before main content
                            if doc_title_cleaned:
                                doc.add_heading("System/Process Flow Diagram", level=1)
                                doc.add_picture(tmp_img.name, width=Pt(600))
                                doc.add_page_break()
                    # Do not insert any fallback heading/code at the end
                except Exception as e:
                    pass  # Do not insert any fallback heading/code at the end

            # Add JSON message structures at the end if any
            if hasattr(st.session_state, "all_stories_data"):
                jsons_to_add = []
                for i, story in enumerate(st.session_state.all_stories_data, 1):
                    for field in ["description", "acceptance_criteria"]:
                        val = story.get(field)
                        if isinstance(val, str):
                            matches = re.findall(r'({[\s\S]+?})', val)
                            for match in matches:
                                try:
                                    parsed = json.loads(match)
                                    jsons_to_add.append((i, story.get('id', 'N/A'), parsed))
                                except Exception:
                                    continue
                if jsons_to_add:
                    doc.add_page_break()
                    doc.add_heading("JSON Message Structures Found in Stories", level=1)
                    for i, story_id, js in jsons_to_add:
                        doc.add_heading(f"Story {i} (ID: {story_id})", level=2)
                        for line in json.dumps(js, indent=2).splitlines():
                            doc.add_paragraph(line, style='Normal')

            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)

            st.download_button(
                label="Download Document as Word File (.docx)",
                data=bio,
                file_name=f"functional_doc_combined_{download_filename_suffix}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Error creating .docx file: {e}")
            st.download_button(
                label="Download Document as Text File (DOCX failed)",
                data=st.session_state.generated_document,
                file_name=f"functional_doc_combined_{download_filename_suffix}.txt",
                mime="text/plain"
            )

    st.markdown("---")
